"""Word doc exporter — generates the official-format schedule docx."""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


CONSULTANT_CODES = {
    "alsayed":"AS","suleman":"KS","akhtar":"SA","alhussaini":"HH",
    "bazarbashi":"SB","anwar":"MA","alghabban":"AGA","alqahtani":"AQ",
    "almugbel":"FA","tweigieri":"TT","atallah":"JA","aljubran":"AA",
    "aljabrani":"AA","alzahrani":"AZ","sabah":"SAK","meshari":"MAZ",
    "rauf":"MR","aisha":"AIS","alyahya":"MAY","chemoassessment":"CA",
}

COLOR_HEADER = "1F4E78"
COLOR_DAY    = "D9E1F2"
COLOR_GREY   = "F2F2F2"
COLOR_CHEMO  = "FFF2CC"
COLOR_ADDON  = "FCE4B5"  # light orange — add-on clinics (consultant on leave)

# Per-role helper-name colors in the schedule grid, per the chief's spec:
#   Fellows -> red, Assistants & NPs -> black, Residents -> green
COLOR_FELLOW    = "C00000"
COLOR_RESIDENT  = "1E7E46"
COLOR_DEFAULT   = "000000"   # NPs, assistants, externals, unrecognized names


def _build_role_colors(result):
    """
    Name (lowercased, no trailing '*') -> hex color, built from the three
    summary tables the solver returns. Fellows/NPs share fellows_summary but
    are distinguished by the Role field; residents and assistants each have
    their own summary list.
    """
    colors = {}
    for f in result.get("fellows_summary", []) or []:
        name = str(f.get("Fellow", "")).strip().lower()
        role = str(f.get("Role", "Fellow")).strip().lower()
        colors[name] = COLOR_DEFAULT if role == "np" else COLOR_FELLOW
    for r in result.get("residents_summary", []) or []:
        name = str(r.get("Resident", "")).strip().lower()
        colors[name] = COLOR_RESIDENT
    for a in result.get("assistants_summary", []) or []:
        name = str(a.get("Assistant", "")).strip().lower()
        colors[name] = COLOR_DEFAULT
    return colors


def _set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _style_run(run, bold=False, size=10, color=None, italic=False):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _write_cell(cell, text, *, bold=False, size=10, color=None,
                bg=None, align="center", italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = p.add_run(text or "")
    _style_run(run, bold=bold, size=size, color=color, italic=italic)
    if bg:
        _set_cell_bg(cell, bg)
    _set_cell_borders(cell)


def _code_for(consultant):
    return CONSULTANT_CODES.get(consultant.strip().lower(), "")


def _format_helpers(s):
    if not s or s.strip() in ("—", "-", ""):
        return ""
    parts = [p.strip() for p in s.split(",") if p.strip()]
    return "/".join(p.upper() for p in parts)


def _write_helpers_cell(cell, assigned_str, role_colors, *, bg=None, size=10):
    """
    Same visual slot as _write_cell but writes each name as its own colored
    run (Fellow=red, Resident=green, NP/Assistant/other=black) joined by '/'.
    """
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    names = [n.strip() for n in (assigned_str or "").split(",") if n.strip()]
    if not names or names == ["—"] or names == ["-"]:
        _set_cell_borders(cell)
        if bg:
            _set_cell_bg(cell, bg)
        return

    for i, raw_name in enumerate(names):
        is_external = raw_name.endswith("*")
        lookup = raw_name.rstrip("*").strip().lower()
        color = role_colors.get(lookup, COLOR_DEFAULT)
        run = p.add_run(raw_name.upper() + ("/" if i < len(names) - 1 else ""))
        _style_run(run, bold=False, size=size, color=color, italic=is_external)
    if bg:
        _set_cell_bg(cell, bg)
    _set_cell_borders(cell)


def export_docx(result, output_path, period_start, period_end, day_info):
    """
    day_info: list of (day_short, day_label, day_date) e.g.
              [("Sun","SUN","10 May"), ("Mon","MON","11 May"), ...]
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin   = Cm(1.2)
        section.right_margin  = Cm(1.2)
        section.orientation   = 1
        section.page_width    = Cm(29.7)
        section.page_height   = Cm(21.0)

    role_colors = _build_role_colors(result)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_run(title.add_run("MEDICAL ONCOLOGY"), bold=True, size=14, color=COLOR_HEADER)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _style_run(sub.add_run("OUTPATIENT AND INPATIENT SCHEDULE"),
               bold=True, size=12, color=COLOR_HEADER)

    ph = doc.add_table(rows=1, cols=4)
    for i, w in enumerate([Cm(3), Cm(8), Cm(8), Cm(8)]):
        ph.columns[i].width = w
    _write_cell(ph.rows[0].cells[0], "PERIOD:", bold=True, bg=COLOR_GREY, align="left")
    _write_cell(ph.rows[0].cells[1], f"{period_start}  -  {period_end}", bold=True)
    _write_cell(ph.rows[0].cells[2], "")
    _write_cell(ph.rows[0].cells[3],
                f"Revised: {datetime.now().strftime('%d %b %Y')}",
                italic=True, align="right")

    doc.add_paragraph()
    h = doc.add_paragraph()
    _style_run(h.add_run("OUTPATIENT COVERAGE"), bold=True, size=12, color=COLOR_HEADER)

    schedule = result["schedule"]
    def clinics_for(day_short, session):
        return [c for c in schedule
                if c["Day"].startswith(day_short) and c["Session"] == session]

    day_row_counts = {ds: max(len(clinics_for(ds, "AM")),
                              len(clinics_for(ds, "PM")), 1)
                      for ds, _, _ in day_info}
    total_rows = 2 + sum(day_row_counts.values())

    table = doc.add_table(rows=total_rows, cols=7)
    for i, w in enumerate([2.2, 2.2, 3.5, 5.5, 2.2, 3.5, 5.5]):
        table.columns[i].width = Cm(w)

    h0 = table.rows[0]
    _write_cell(h0.cells[0], "Day &", bold=True, color="FFFFFF", bg=COLOR_HEADER)
    am = h0.cells[1].merge(h0.cells[2]).merge(h0.cells[3])
    _write_cell(am, "MORNING CLINICS", bold=True, color="FFFFFF", bg=COLOR_HEADER)
    pm = h0.cells[4].merge(h0.cells[5]).merge(h0.cells[6])
    _write_cell(pm, "AFTERNOON CLINICS", bold=True, color="FFFFFF", bg=COLOR_HEADER)

    h1 = table.rows[1]
    for i, lab in enumerate(["Date","Clinic Code","Consultant","Asst Cons │ Fellow/Res",
                             "Clinic Code","Consultant","Asst Cons │ Fellow/Res"]):
        _write_cell(h1.cells[i], lab, bold=True, color="FFFFFF", bg=COLOR_HEADER, size=9)

    row_idx = 2
    for day_short, day_label, day_date in day_info:
        am_c = clinics_for(day_short, "AM")
        pm_c = clinics_for(day_short, "PM")
        n_rows = day_row_counts[day_short]

        if n_rows > 1:
            merged = table.rows[row_idx].cells[0].merge(
                     table.rows[row_idx + n_rows - 1].cells[0])
            _write_cell(merged, f"{day_label}\n{day_date}", bold=True, bg=COLOR_DAY, size=11)
        else:
            _write_cell(table.rows[row_idx].cells[0],
                        f"{day_label}\n{day_date}", bold=True, bg=COLOR_DAY, size=11)

        for i in range(n_rows):
            cells = table.rows[row_idx + i].cells
            if i < len(am_c):
                c = am_c[i]
                is_chemo = c["IsChemo"]
                is_addon = c.get("IsAddOn", False)
                pts = c.get("Patients")
                pts_suffix = f" ({pts})" if (pts not in (None, "") and not is_chemo) else ""
                if is_chemo:
                    bg = COLOR_CHEMO
                    label = "CHEMOASSESSMENT"
                elif is_addon:
                    bg = COLOR_ADDON
                    label = f"(ADD-ON) {c['Consultant'].upper()}{pts_suffix}"
                else:
                    bg = None
                    label = f"{c['Consultant'].upper()}{pts_suffix}"
                _write_cell(cells[1], _code_for(c["Consultant"]), bold=True, bg=bg)
                _write_cell(cells[2], label, bold=(is_chemo or is_addon), bg=bg)
                _write_helpers_cell(cells[3], c["Assigned"], role_colors, bg=bg)
            else:
                for col in (1, 2, 3):
                    _write_cell(cells[col], "")
            if i < len(pm_c):
                c = pm_c[i]
                is_chemo = c["IsChemo"]
                is_addon = c.get("IsAddOn", False)
                pts = c.get("Patients")
                pts_suffix = f" ({pts})" if (pts not in (None, "") and not is_chemo) else ""
                if is_chemo:
                    bg = COLOR_CHEMO
                    label = "CHEMOASSESSMENT"
                elif is_addon:
                    bg = COLOR_ADDON
                    label = f"(ADD-ON) {c['Consultant'].upper()}{pts_suffix}"
                else:
                    bg = None
                    label = f"{c['Consultant'].upper()}{pts_suffix}"
                _write_cell(cells[4], _code_for(c["Consultant"]), bold=True, bg=bg)
                _write_cell(cells[5], label, bold=(is_chemo or is_addon), bg=bg)
                _write_helpers_cell(cells[6], c["Assigned"], role_colors, bg=bg)
            else:
                for col in (4, 5, 6):
                    _write_cell(cells[col], "")
        row_idx += n_rows

    doc.add_paragraph()
    leg = doc.add_table(rows=2, cols=4)
    for i in range(4):
        leg.columns[i].width = Cm(7)
    _write_cell(leg.rows[0].cells[0], "LEGEND:", bold=True, bg=COLOR_GREY, size=9)
    _write_cell(leg.rows[0].cells[1], "* = External staff (purple rule)", size=9, align="left")
    _write_cell(leg.rows[0].cells[2], "CA = Chemoassessment Clinic", size=9, align="left")
    _write_cell(leg.rows[0].cells[3], "CB = Combined Clinic (varies weekly)", size=9, align="left")
    _write_cell(leg.rows[1].cells[0], "", bg=COLOR_GREY, size=9)
    _write_cell(leg.rows[1].cells[1], "(ADD-ON) = consultant on leave; covered solo", size=9, align="left")
    _write_cell(leg.rows[1].cells[2], "", size=9, align="left")
    _write_cell(leg.rows[1].cells[3], "", size=9, align="left")

    color_note = doc.add_paragraph()
    r1 = color_note.add_run("Color code: ")
    _style_run(r1, bold=True, size=9)
    r2 = color_note.add_run("Fellows")
    _style_run(r2, bold=True, size=9, color=COLOR_FELLOW)
    r3 = color_note.add_run("   Residents")
    _style_run(r3, bold=True, size=9, color=COLOR_RESIDENT)
    r4 = color_note.add_run("   NPs / Assistants")
    _style_run(r4, bold=True, size=9, color=COLOR_DEFAULT)

    # Workload summary tables
    doc.add_paragraph()
    h = doc.add_paragraph()
    _style_run(h.add_run("WORKLOAD SUMMARY"), bold=True, size=12, color=COLOR_HEADER)

    sub = doc.add_paragraph()
    _style_run(sub.add_run("Fellows"), bold=True, size=10, color=COLOR_HEADER)
    fm = result["fellows_summary"]
    ft = doc.add_table(rows=len(fm) + 1, cols=6)
    for i, w in enumerate([4.0, 3.5, 2.0, 2.5, 3.0, 3.0]):
        ft.columns[i].width = Cm(w)
    for i, lab in enumerate(["Fellow","Rotation","Role","Total","Inside","Rotation %"]):
        _write_cell(ft.rows[0].cells[i], lab, bold=True, color="FFFFFF", bg=COLOR_HEADER)
    for i, f in enumerate(fm, start=1):
        _write_cell(ft.rows[i].cells[0], f["Fellow"], align="left")
        _write_cell(ft.rows[i].cells[1], f["Rotation"])
        _write_cell(ft.rows[i].cells[2], f.get("Role", "Fellow"))
        _write_cell(ft.rows[i].cells[3], str(f["Total"]), bold=True)
        _write_cell(ft.rows[i].cells[4], str(f["Inside"]))
        pct = f"{round(100*f['Inside']/f['Total'])}%" if f["Total"] else "—"
        _write_cell(ft.rows[i].cells[5], pct)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    _style_run(sub.add_run("Residents"), bold=True, size=10, color=COLOR_HEADER)
    rm = result["residents_summary"]
    rt = doc.add_table(rows=len(rm) + 1, cols=2)
    for i, w in enumerate([6, 4]):
        rt.columns[i].width = Cm(w)
    for i, lab in enumerate(["Resident","Total Clinics"]):
        _write_cell(rt.rows[0].cells[i], lab, bold=True, color="FFFFFF", bg=COLOR_HEADER)
    for i, r in enumerate(rm, start=1):
        _write_cell(rt.rows[i].cells[0], r["Resident"], align="left")
        _write_cell(rt.rows[i].cells[1], str(r["Total"]), bold=True)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
